"""
NEXO SOBERANO — Video Analysis & DOCX Export
Endpoints:
  POST /api/agente/analizar-video   — transcribe + analiza video (archivo o YouTube URL)
  POST /api/agente/exportar-docx    — genera .docx con transcripción + análisis
"""

import os
import re
import shutil
import logging
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from pydantic import BaseModel

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/agente", tags=["video"])


# ─── Helpers ────────────────────────────────────────────────────────────────

DOMAIN_HINTS = {
    "geo":      "Geopolítica — conflictos, alianzas, poder regional",
    "conducta": "Conducta de líderes — lenguaje corporal, patrones, decisiones",
    "masas":    "Control de masas — narrativas, expectativa social, psicología",
    "eco_aus":  "Economía Austriaca — ciclos, capital, anarcocapitalismo",
    "poder":    "Escenarios de poder — actores, vectores, tablero político",
    "osint":    "OSINT — fuentes abiertas, señales, patrones",
    "libre":    "Análisis general sin restricción de dominio",
}

ANALYSIS_PROMPT = """Eres el sistema de inteligencia estratégica NEXO SOBERANO.
Dominio de análisis: {hint}

Se te proporciona la transcripción completa de un video titulado: "{titulo}"
Canal/Fuente: {canal}

Genera un análisis de inteligencia estructurado con exactamente estas secciones:
1. RESUMEN EJECUTIVO
2. PUNTOS CLAVE IDENTIFICADOS
3. ANÁLISIS ESTRATÉGICO ({dominio_upper})
4. SEÑALES DÉBILES Y SUBTEXT
5. CONCLUSIONES E IMPLICACIONES
6. VERIFICACIÓN RECOMENDADA

Transcripción:
---
{transcripcion}
---

Responde en español. Sé preciso y denso en información, sin relleno."""


# ─── POST /api/agente/analizar-video ────────────────────────────────────────

@router.post("/analizar-video")
async def analizar_video(
    file: Optional[UploadFile] = File(default=None),
    youtube_url: Optional[str] = Form(default=None),
    dominio: Optional[str] = Form(default="libre"),
    idioma: Optional[str] = Form(default="es"),
):
    """
    Transcribe un video (archivo subido o URL de YouTube) con OpenAI Whisper
    y lo analiza con el sistema de inteligencia NEXO.
    """
    import openai

    OPENAI_KEY = os.getenv("OPENAI_API_KEY", "")
    if not OPENAI_KEY:
        raise HTTPException(
            status_code=503,
            detail="OPENAI_API_KEY no configurada — necesaria para transcripcion Whisper"
        )

    tmp_dir = Path(tempfile.mkdtemp())
    audio_path: Optional[Path] = None
    titulo_video = "Video sin titulo"
    canal = "Desconocido"
    duracion_seg: float = 0

    try:
        # ── 1. Obtener audio ──────────────────────────────────────────────────
        if youtube_url:
            try:
                import yt_dlp
            except ImportError:
                raise HTTPException(
                    status_code=503,
                    detail="yt-dlp no instalado. Ejecuta: pip install yt-dlp"
                )

            ydl_opts = {
                "format": "bestaudio/best",
                "outtmpl": str(tmp_dir / "audio.%(ext)s"),
                "postprocessors": [{
                    "key": "FFmpegExtractAudio",
                    "preferredcodec": "mp3",
                    "preferredquality": "128",
                }],
                "quiet": True,
                "no_warnings": True,
            }
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(youtube_url, download=True)
                titulo_video = info.get("title", titulo_video)
                duracion_seg = float(info.get("duration", 0))
                canal = info.get("uploader", canal)

            # yt-dlp puede nombrar el archivo de otra forma
            audio_candidates = list(tmp_dir.glob("audio.*"))
            if not audio_candidates:
                raise HTTPException(status_code=500, detail="No se pudo descargar el audio de YouTube")
            audio_path = audio_candidates[0]

        elif file and file.filename:
            ext = Path(file.filename).suffix.lower()
            raw_path = tmp_dir / f"input{ext}"
            with open(raw_path, "wb") as f:
                content = await file.read()
                f.write(content)
            titulo_video = file.filename
            canal = "Archivo local"

            # Extraer audio con ffmpeg
            mp3_path = tmp_dir / "audio.mp3"
            result = subprocess.run(
                ["ffmpeg", "-i", str(raw_path), "-vn", "-acodec", "mp3",
                 "-ab", "128k", "-y", str(mp3_path)],
                capture_output=True, timeout=300
            )
            audio_path = mp3_path if (result.returncode == 0 and mp3_path.exists()) else raw_path
        else:
            raise HTTPException(
                status_code=400,
                detail="Se requiere un archivo (file) o una URL de YouTube (youtube_url)"
            )

        # ── 2. Transcribir con Whisper ────────────────────────────────────────
        client = openai.OpenAI(api_key=OPENAI_KEY)
        with open(audio_path, "rb") as af:
            transcript_resp = client.audio.transcriptions.create(
                model="whisper-1",
                file=af,
                language=idioma,
                response_format="verbose_json",
            )

        transcripcion: str = transcript_resp.text or ""
        duracion_whisper: float = float(getattr(transcript_resp, "duration", 0) or duracion_seg)

        # ── 3. Analizar con NEXO ──────────────────────────────────────────────
        hint = DOMAIN_HINTS.get(dominio, DOMAIN_HINTS["libre"])
        prompt = ANALYSIS_PROMPT.format(
            hint=hint,
            titulo=titulo_video,
            canal=canal,
            dominio_upper=(dominio or "libre").upper(),
            transcripcion=transcripcion[:12000],
        )

        analisis = ""
        try:
            # Intentar usar ai_router del proyecto
            from NEXO_CORE.services.ai_router import ai_router, AIRequest
            ai_resp = await ai_router.consultar(AIRequest(
                pregunta=prompt,
                modo=dominio or "libre",
                usar_rag=False,
                max_tokens=2000,
            ))
            analisis = ai_resp.get("respuesta") or ai_resp.get("answer") or ""
        except Exception as e_ai:
            logger.warning(f"ai_router no disponible, usando openai directo: {e_ai}")
            # Fallback: Claude via openai compatible o GPT-4o
            try:
                chat_resp = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=2000,
                )
                analisis = chat_resp.choices[0].message.content or ""
            except Exception as e2:
                analisis = f"[Analisis automatico no disponible: {e2}]"

        return {
            "ok": True,
            "titulo": titulo_video,
            "canal": canal,
            "duracion_seg": duracion_whisper or duracion_seg,
            "dominio": dominio,
            "transcripcion": transcripcion,
            "analisis": analisis,
            "fecha": datetime.now(timezone.utc).isoformat(),
            "palabras": len(transcripcion.split()),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en analizar-video: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ─── POST /api/agente/exportar-docx ─────────────────────────────────────────

class DocxRequest(BaseModel):
    titulo: str = "Analisis NEXO SOBERANO"
    canal: Optional[str] = ""
    duracion_seg: Optional[float] = 0
    dominio: Optional[str] = "libre"
    fecha: Optional[str] = ""
    transcripcion: str
    analisis: str
    palabras: Optional[int] = 0


@router.post("/exportar-docx")
async def exportar_docx(req: DocxRequest):
    """
    Genera un .docx con transcripcion + analisis de video.
    Retorna el archivo para descarga directa.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=503, detail="python-docx no instalado. Ejecuta: pip install python-docx")

    doc = Document()

    # Estilo base
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ── Encabezado ────────────────────────────────────────────────────────────
    h0 = doc.add_heading("", level=0)
    r0 = h0.add_run("NEXO SOBERANO")
    r0.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
    r0.font.size = Pt(22)
    r0.bold = True

    h1 = doc.add_heading("", level=2)
    r1 = h1.add_run("ANALISIS DE INTELIGENCIA — VIDEO")
    r1.font.color.rgb = RGBColor(0x22, 0x55, 0x22)
    r1.font.size = Pt(13)

    doc.add_paragraph()

    # ── Tabla metadata ────────────────────────────────────────────────────────
    fecha_str = (req.fecha or "")[:10] or datetime.now().strftime("%Y-%m-%d")
    dur_m = int((req.duracion_seg or 0) // 60)
    dur_s = int((req.duracion_seg or 0) % 60)
    dur_str = f"{dur_m}:{dur_s:02d}" if req.duracion_seg else "—"

    meta = [
        ("Titulo", req.titulo),
        ("Canal / Fuente", req.canal or "—"),
        ("Duracion", dur_str),
        ("Dominio de Analisis", (req.dominio or "libre").upper()),
        ("Fecha de Analisis", fecha_str),
        ("Palabras transcritas", str(req.palabras or len(req.transcripcion.split()))),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate(meta):
        tbl.rows[i].cells[0].text = lbl
        run_lbl = tbl.rows[i].cells[0].paragraphs[0].runs
        if run_lbl:
            run_lbl[0].bold = True
        tbl.rows[i].cells[1].text = val

    doc.add_paragraph()
    doc.add_paragraph("=" * 60)
    doc.add_paragraph()

    # ── Seccion: Analisis ─────────────────────────────────────────────────────
    h_analisis = doc.add_heading("ANALISIS DE INTELIGENCIA", level=1)
    h_analisis.runs[0].font.color.rgb = RGBColor(0x1A, 0x5C, 0x2A)
    doc.add_paragraph()

    for line in req.analisis.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        # Secciones numeradas: "1. TITULO"
        if line and len(line) > 2 and line[0].isdigit() and line[1] in (".", ")"):
            h = doc.add_heading("", level=2)
            run_h = h.add_run(line)
            run_h.font.color.rgb = RGBColor(0x1A, 0x5C, 0x2A)
        elif line.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])
        else:
            doc.add_paragraph(line)

    doc.add_paragraph()
    doc.add_paragraph("=" * 60)
    doc.add_paragraph()

    # ── Seccion: Transcripcion ────────────────────────────────────────────────
    h_trans = doc.add_heading("TRANSCRIPCION COMPLETA", level=1)
    h_trans.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x6C)
    doc.add_paragraph(f"Total de palabras: {req.palabras or len(req.transcripcion.split())}")
    doc.add_paragraph()

    # Partir en parrafos de ~500 chars respetando espacios
    texto = req.transcripcion.strip()
    chunk_size = 500
    while texto:
        end = min(chunk_size, len(texto))
        if len(texto) > chunk_size:
            sp = texto[:chunk_size].rfind(" ")
            end = sp if sp > 200 else chunk_size
        p = doc.add_paragraph(texto[:end].strip())
        p.paragraph_format.space_after = Pt(4)
        texto = texto[end:].strip()

    # ── Pie de documento ──────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph("=" * 60)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        f"Generado por NEXO SOBERANO  |  {fecha_str}  |  Clasificacion: USO INTERNO"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Guardar y devolver ────────────────────────────────────────────────────
    safe = re.sub(r"[^\w\s-]", "", req.titulo)[:40].strip().replace(" ", "_")
    filename = f"NEXO_{safe}_{fecha_str}.docx"
    out_path = Path(tempfile.mkdtemp()) / filename
    doc.save(str(out_path))

    return FileResponse(
        path=str(out_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
