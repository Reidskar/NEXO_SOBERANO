import os
import sqlite3
import json
from datetime import datetime

# try to import Document from python-docx; provide guidance if missing
try:
    from docx import Document
except ImportError:
    raise ImportError(
        "\nModule 'docx' not found.\nPlease install dependencies by running:\n    python -m pip install -r requirements.txt\n"      
    )

BASE_PATH = "NEXO_SOBERANO"

def crear_infraestructura():
    # Estructura modular para escalabilidad de Terabytes
    carpetas = [
        "boveda_datos",           # Registro de archivos 6TB
        "memoria_vectorial",      # ChromaDB (Pensamiento semántico)
        "base_sqlite",            # Cerebro de metadatos
        "agentes_autonomos",      # Scripts de fondo
        "modulos_experimentales", # IA en pruebas
        "bitacora",               # Triple registro (MD, JSON, DOCX)
        "web_api"                 # Puente hacia el público
    ]
    
    for carpeta in carpetas:
        ruta = os.path.join(BASE_PATH, carpeta)
        os.makedirs(ruta, exist_ok=True)
        marker = os.path.join(ruta, "README.md")
        if not os.path.exists(marker):
            with open(marker, "w", encoding="utf-8") as f:
                f.write(f"Carpeta de trabajo: {carpeta}\n")
    log.info("✅ Estructura de carpetas creada.")

def crear_boveda_db():
    db_path = os.path.join(BASE_PATH, "base_sqlite", "boveda.db")
    conn = sqlite3.connect(db_path)
    # Tabla de Evidencia con Trazabilidad Total
    conn.execute("""
    CREATE TABLE IF NOT EXISTS evidencia (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        hash_sha256 TEXT UNIQUE,
        nombre_archivo TEXT,
        ruta_local TEXT,
        link_nube TEXT,
        dominio TEXT,        -- Local, Drive, OneDrive
        categoria TEXT,
        resumen_ia TEXT,
        fecha_ingesta TIMESTAMP,
        nivel_confianza REAL,
        impacto TEXT,
        vectorizado INTEGER DEFAULT 0
    )
    """)
    conn.commit()
    conn.close()
    log.info("✅ Bóveda SQLite inicializada.")

def inicializar_bitacoras():
    fecha = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # 1. Bitácora Markdown (Para lectura rápida)
    md_path = os.path.join(BASE_PATH, "bitacora", "evolucion.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"# NEXO SOBERANO - Bitácora de Evolución\n\n## [v1.0.0] - {fecha}\n- Sistema fundado.")

    # 2. Bitácora JSON (Para que la IA se auto-analice)
    json_path = os.path.join(BASE_PATH, "bitacora", "evolucion.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump({"entradas": [{"fecha": fecha, "evento": "Inicialización", "v": "1.0.0"}]}, f)

    # 3. Bitácora DOCX (Respaldo institucional/público)
    doc = Document()
    doc.add_heading('Bitácora de Inteligencia Nexo Soberano', 0)
    doc.add_paragraph(f'Inicio de operaciones: {fecha}')
    doc.save(os.path.join(BASE_PATH, "bitacora", "evolucion.docx"))
    log.info("✅ Triple bitácora establecida.")

if __name__ == "__main__":
    crear_infraestructura()
    crear_boveda_db()
    inicializar_bitacoras()
    log.info("\n🚀 SISTEMA NEXO SOBERANO NIVEL 1 ACTIVADO.")
